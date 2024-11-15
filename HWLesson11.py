# Задание №1 
# а) Прочитайте из трёх excel файлов (заранее создайте их, внутри 1111, 2222, 3333). 

import openpyxl
from openpyxl import Workbook
from openpyxl.styles import Font, Border, Side

wb1 = Workbook()
wb2 = Workbook()
wb3 = Workbook()

sheet1 = wb1.active
sheet2 = wb2.active
sheet3 = wb3.active

sheet1.cell(row = 1, column = 1, value = 1111)
sheet2.cell(row = 1, column = 1, value = 2222)
sheet3.cell(row = 1, column = 1, value = 3333)

wb1.save("file1.xlsx")
print("Файл 'file1.xlsx' создан.")
wb2.save("file2.xlsx")
print("Файл 'file2.xlsx' создан.")
wb3.save("file3.xlsx")
print("Файл 'file3.xlsx' создан.")

def readDataFromExcel(filePath):
    wb = openpyxl.load_workbook(filePath)
    sheet = wb.active
    data = []
    for row in sheet.iter_rows(values_only = True):
        data.extend(row)
    return data

data1 = readDataFromExcel("file1.xlsx")
data2 = readDataFromExcel("file2.xlsx")
data3 = readDataFromExcel("file3.xlsx")

print("file1.xlsx: ", data1)
print("file2.xlsx: ", data2)
print("file3.xlsx: ", data3)

# б) Отсортируйте полученную матрицу в порядке убывания. 
fullData = data1 + data2 + data3
sortedData = sorted(fullData, reverse = True)

# в) Запишите это в один файл, изменив шрифт и обернув в границы.
wb = openpyxl.Workbook()
sheet = wb.active

boldFont = Font(bold = True)
thinBorder = Border(left = Side(style = 'thin'), right = Side(style = 'thin'), top=Side(style = 'thin'), bottom = Side(style = 'thin'))

for i, value in enumerate(sortedData, start=1):
    cell = sheet.cell(row = 1, column = i, value = value)
    cell.font = boldFont
    cell.border = thinBorder

wb.save("sorted_data.xlsx")
print("Данные успешно записаны в файл 'sorted_data.xlsx'.")

# а) Создайте json файл в операционной системе, заполните его данными с сайта 
# https://jsonplaceholder.typicode.com/todos/ 

import requests
import json

url = "https://jsonplaceholder.typicode.com/todos/"

resp = requests.get(url)
data = resp.json()

with open("data.json", "w") as file:
    json.dump(data, file, indent = 4)
print("Данные сохранены в data.json")

# б) Прочитайте этот файл в массив python-словарей. 
import json

with open("data.json", "r") as file:
    data = json.load(file)

print(f"Data type: {type(data)}")

# в) Запишите каждый из этих словарей в отдельный json-файл.
import json

with open("data.json", "r") as file:
    data = json.load(file)

for index, item in enumerate(data, start = 1):
    file_name = f"data_{index}.json"
    with open(file_name, "w") as json_file:
        json.dump(item, json_file, indent=4)
    print(f"Saved {file_name}")

print("Все словари были сохранены в файлы JSON.")

# а) Создайте word файл в операционной системе, заполните его текстом «Hello Python». 

from docx import Document

doc = Document()
doc.add_paragraph("Hello Python")
doc.save("python1.docx")

print("Файл успешно создан")

# б) Прочитайте этот файл, только жирный текст в python-строку и выведите на экран. 
from docx import Document

doc = Document("python1.docx")
txt = "\n".join([paragraph.text for paragraph in doc.paragraphs])

print("Текст из файла:")
print(txt)

# в) Создайте абзац с текстом и запишите это в новый word-файл, изменит шрифт и размер 
# шрифта.

from docx import Document
doc = Document()
doc.add_paragraph("Вновь созданный файл с записанным текстом.")

doc.save("python2.docx")

print("Файл python2.docx создан.")